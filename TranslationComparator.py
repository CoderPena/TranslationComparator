from docx import Document
from deep_translator import GoogleTranslator
from difflib import SequenceMatcher
from tqdm import tqdm  # Biblioteca para barra de progresso

# Defina o limiar de similaridade como variável
SIMILARITY_THRESHOLD = 0.9  # Limiar de similaridade

# Caminhos dos arquivos
file_path = "/home/pena/Documents/TRPF - Tradução/TRPF026/"
pt_file_name = "PROPOSTA_EAD.docx"
en_file_name = "PROPOSTA_EAD PTBR-EN.docx"
df_file_name = "PROPOSTA_EAD PTBR-EN_DIFF.docx"

pt_file = file_path + pt_file_name
en_file = file_path + en_file_name
# Gerar relatório
diff_report_path = file_path + "comparacao_traducao_relatorio.txt"
en_file_diff = file_path + df_file_name

# Função para carregar texto de um arquivo DOCX com barra de progresso
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = []
    table_metadata = []  # Para armazenar informações sobre tabelas

    # Processar parágrafos com barra de progresso
    paragraphs = doc.paragraphs
    print(f"Lendo parágrafos do arquivo: {file_path}")
    for i, paragraph in enumerate(tqdm(paragraphs, desc="Parágrafos processados")):
        if paragraph.text.strip():  # Ignorar linhas vazias
            text.append(("parágrafo", i + 1, paragraph.text.strip()))

    # Processar tabelas com barra de progresso
    tables = doc.tables
    print(f"Lendo tabelas do arquivo: {file_path}")
    for table_index, table in enumerate(tqdm(tables, desc="Tabelas processadas")):
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                table_metadata.append((table_index + 1, row_index + 1, col_index + 1, bool(cell_text)))
                if cell_text:  # Ignorar células vazias
                    text.append(("tabela", table_index + 1, row_index + 1, col_index + 1, cell_text))

    return text, table_metadata

# Função para contar palavras em uma lista de texto
def count_words(text_list):
    return sum(len(entry[-1].split()) for entry in text_list)

# Função para comparar textos com barra de progresso
def compare_texts(pt_text, en_text):
    differences = []
    total_comparisons = 0
    print("Comparando textos...")
    for i, (pt_entry, en_entry) in enumerate(tqdm(zip(pt_text, en_text), total=min(len(pt_text), len(en_text)), desc="Comparações realizadas")):
        pt_type, *pt_metadata, pt_line = pt_entry
        en_type, *en_metadata, en_line = en_entry
        translated_line = GoogleTranslator(source='portuguese', target='english').translate(pt_line)
        similarity = SequenceMatcher(None, translated_line, en_line).ratio()
        total_comparisons += 1
        if similarity < SIMILARITY_THRESHOLD:  # Usando o limiar configurável
            differences.append((len(differences) + 1, i + 1, pt_entry, en_entry, translated_line, similarity))
    return differences, total_comparisons

# Carregar os textos dos arquivos
pt_text, pt_table_metadata = extract_text_from_docx(pt_file)
en_text, en_table_metadata = extract_text_from_docx(en_file)

# Contar palavras nos textos
pt_word_count = count_words(pt_text)
en_word_count = count_words(en_text)

# Comparar os textos
differences, total_comparisons = compare_texts(pt_text, en_text)

doc_en = Document(en_file)  # Abrir o arquivo de destino para modificação

with open(diff_report_path, "w", encoding="utf-8") as report:
    report.write(f"Limiar de similaridade aplicado: {SIMILARITY_THRESHOLD:.2%}\n\n")

    report.write(f"Arquivo em Português: {pt_file}\n")
    report.write(f"Palavras lidas: {pt_word_count}\n\n")
    report.write(f"Arquivo em Inglês: {en_file}\n")
    report.write(f"Palavras lidas: {en_word_count}\n\n")

    # Resumo de tabelas
    total_cells = len(pt_table_metadata)
    filled_cells = sum(1 for _, _, _, filled in pt_table_metadata if filled)
    empty_cells = total_cells - filled_cells
    report.write(f"Total de células em tabelas: {total_cells}\n")
    report.write(f"Células preenchidas: {filled_cells}\n")
    report.write(f"Células vazias: {empty_cells}\n\n")

    report.write(f"Total de comparações realizadas: {total_comparisons}\n")
    report.write(f"Total de diferenças encontradas: {len(differences)}\n\n")

    if differences:
        suggestion_number = 1  # Contador de sugestões
        report.write("Diferenças encontradas:\n\n")
        for diff_number, diff_index, pt_entry, en_entry, translated_line, similarity in differences:
            pt_type, *pt_metadata, pt_text = pt_entry
            en_type, *en_metadata, en_text = en_entry

            report.write(f"Número da Diferença: {diff_number}\n")
            if pt_type == "parágrafo":
                report.write(f"Tipo: Parágrafo\nNúmero da linha no documento: {pt_metadata[0]}\n")
            else:
                report.write(f"Tipo: Tabela\nNúmero da tabela: {pt_metadata[0]}\nLinha da tabela: {pt_metadata[1]}\nColuna da tabela: {pt_metadata[2]}\n")

            report.write(f"Número da Célula: {diff_index}\n")
            report.write(f"PT ORIGINAL: {pt_text}\n")
            report.write(f"EN ORIGINAL: {en_text}\n")
            report.write(f"TRAD SUGERIDA: {translated_line}\n")
            report.write(f"Similaridade: {similarity:.2%}\n")
            report.write("-" * 50 + "\n")

            # Modificar o arquivo de inglês para adicionar sugestão na célula correspondente
            if en_type == "tabela":
                # Encontrar a célula correspondente na tabela e adicionar a sugestão de correção
                table_index, row_index, col_index = en_metadata
                table = doc_en.tables[table_index - 1]  # Índice baseado em 0
                cell = table.rows[row_index - 1].cells[col_index - 1]

                # Adicionar o texto original e a sugestão de correção dentro da mesma célula
                cell.text = f"\n\n[SUGESTÃO] #{suggestion_number}\n------->PT ORIGINAL\n{pt_text}\n------->EN ORIGINAL\n{en_text}\n------->TRAD SUGERIDA\n{translated_line}\n"
                suggestion_number += 1
            elif en_type == "parágrafo":
                # Modificar o parágrafo correspondente e adicionar a sugestão
                paragraph_index = en_metadata[0] - 1  # Índice baseado em 0
                paragraph = doc_en.paragraphs[paragraph_index]

                # Adicionar o texto original e a sugestão de correção dentro do parágrafo
                paragraph.text = f"\n\n[SUGESTÃO] #{suggestion_number}\n------->PT ORIGINAL\n{pt_text}\n------->EN ORIGINAL\n{paragraph.text}\n------->TRAD SUGERIDA\n{translated_line}\n"
                suggestion_number += 1

    else:
        report.write("Os textos estão alinhados e semelhantes!\n")

    # Salvar o arquivo modificado com sufixo DIFF
    doc_en.save(en_file_diff)

print(f"Relatório de diferenças salvo em: {diff_report_path}")
print(f"Arquivo DIFF salvo em: {en_file_diff}")

# Exibição final no console
print("\nResumo do processamento:")
print(f"Arquivo em Português: {pt_file}")
print(f"Palavras lidas: {pt_word_count}")

print(f"\nArquivo em Inglês: {en_file}")
print(f"Palavras lidas: {en_word_count}")

print(f"\nTotal de comparações realizadas: {total_comparisons}")
print(f"Total de diferenças encontradas: {len(differences)}")
print("\nProcessamento concluído.")

