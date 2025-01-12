from docx import Document
from deep_translator import GoogleTranslator
from difflib import SequenceMatcher
from tqdm import tqdm  # Biblioteca para barra de progresso

# Carregar o arquivo original e o novo documento de saída
doc = Document(en_file)
diff_doc = Document()

suggestion_number = 1  # Contador de sugestões

# Iterar pelas entradas e comparar o texto
for pt_entry, en_entry in zip(pt_text, en_text):
    if len(pt_entry) < 3 or len(en_entry) < 3:
        continue  # Ignorar entradas com dados insuficientes

    pt_type, *pt_metadata, pt_line = pt_entry
    en_type, *en_metadata, en_line = en_entry
    translated_line = GoogleTranslator(source='portuguese', target='english').translate(pt_line)
    similarity = SequenceMatcher(None, translated_line, en_line).ratio()

    # Verificar se a similaridade é menor que o limiar
    if similarity < similarity_threshold:
        if en_type == "tabela":
            # Caso a entrada seja uma tabela, adicione as sugestões nela
            table_index, row_index, col_index = en_metadata
            table = doc.tables[table_index - 1]
            cell = table.cell(row_index - 1, col_index - 1)
            cell_text = cell.text.strip()
            new_text = f"{cell_text}\n\n[SUGESTÃO {suggestion_number}] {translated_line}"
            cell.text = new_text
            suggestion_number += 1
        else:
            # Caso contrário, adicione os parágrafos com sugestões
            para = diff_doc.add_paragraph()
            para.add_run(f"Texto Original: {en_line}\n")
            para.add_run(f"[SUGESTÃO {suggestion_number}] {translated_line}\n")
            suggestion_number += 1

# Salvar o arquivo com as sugestões
diff_doc.save(diff_en_file_path)
print(f"Arquivo com as sugestões salvo em: {diff_en_file_path}")
