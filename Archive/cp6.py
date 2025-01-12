from docx import Document
from deep_translator import GoogleTranslator
from difflib import SequenceMatcher
from tqdm import tqdm  # Biblioteca para barra de progresso

# Função para carregar texto de um arquivo DOCX com barra de progresso
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = []
    table_metadata = []  # Para armazenar informações sobre tabelas

    # Processar parágrafos com barra de progresso
    paragraphs = doc.paragraphs
    print(f"Lendo parágrafos do arquivo: {file_path}")
    for i, paragraph in enumerate(tqdm(paragraphs, desc="Parágrafos processados")):
        if paragraph.text.strip():  # Ignorar linhas vazias
            text.append(("parágrafo", i + 1, paragraph.text.strip()))

    # Processar tabelas com barra de progresso
    tables = doc.tables
    print(f"Lendo tabelas do arquivo: {file_path}")
    for table_index, table in enumerate(tqdm(tables, desc="Tabelas processadas")):
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                table_metadata.append((table_index + 1, row_index + 1, col_index + 1, bool(cell_text)))
                if cell_text:  # Ignorar células vazias
                    text.append(("tabela", table_index + 1, row_index + 1, col_index + 1, cell_text))

    return text, table_metadata

# Função para contar palavras em uma lista de texto
def count_words(text_list):
    return sum(len(entry[-1].split()) for entry in text_list)

# Função para comparar textos com barra de progresso
def compare_texts(pt_text, en_text):
    differences = []
    total_comparisons = 0
    print("Comparando textos...")
    for i, (pt_entry, en_entry) in enumerate(tqdm(zip(pt_text, en_text), total=min(len(pt_text), len(en_text)), desc="Comparações realizadas")):
        pt_type, *pt_metadata, pt_line = pt_entry
        en_type, *en_metadata, en_line = en_entry
        translated_line = GoogleTranslator(source='portuguese', target='english').translate(pt_line)
        similarity = SequenceMatcher(None, translated_line, en_line).ratio()
        total_comparisons += 1
        if similarity < 0.9:  # Ajuste o limiar de similaridade
            differences.append((i + 1, pt_entry, en_entry, translated_line, similarity))
    return differences, total_comparisons

# Caminhos dos arquivos
pt_file = "/home/pena/Downloads/Tema_9_-_ESPERANÇA2024-12-06_16_56_37-1.docx"
en_file = "/home/pena/Downloads/Tema_9_-_ESPERANÇA2024-12-06_16_56_37-1 PTBR-EN.docx"

# Carregar os textos dos arquivos
pt_text, pt_table_metadata = extract_text_from_docx(pt_file)
en_text, en_table_metadata = extract_text_from_docx(en_file)

# Contar palavras nos textos
pt_word_count = count_words(pt_text)
en_word_count = count_words(en_text)

# Comparar os textos
differences, total_comparisons = compare_texts(pt_text, en_text)

# Gerar relatório
diff_report_path = "/home/pena/Downloads/comparacao_traducao_relatorio.txt"
with open(diff_report_path, "w", encoding="utf-8") as report:
    accuracy_threshold = 0.9  # Limiar de similaridade usado
    report.write(f"Limiar de similaridade aplicado: {accuracy_threshold:.2%}\n\n")

    report.write(f"Arquivo em Português: {pt_file}\n")
    report.write(f"Palavras lidas: {pt_word_count}\n\n")
    report.write(f"Arquivo em Inglês: {en_file}\n")
    report.write(f"Palavras lidas: {en_word_count}\n\n")

    # Resumo de tabelas
    total_cells = len(pt_table_metadata)
    filled_cells = sum(1 for _, _, _, filled in pt_table_metadata if filled)
    empty_cells = total_cells - filled_cells
    report.write(f"Total de células em tabelas: {total_cells}\n")
    report.write(f"Células preenchidas: {filled_cells}\n")
    report.write(f"Células vazias: {empty_cells}\n\n")

    report.write(f"Total de comparações realizadas: {total_comparisons}\n")
    report.write(f"Total de diferenças encontradas: {len(differences)}\n\n")

    if differences:
        report.write("Diferenças encontradas:\n\n")
        for diff_index, pt_entry, en_entry, translated_line, similarity in differences:
            pt_type, *pt_metadata, pt_text = pt_entry
            en_type, *en_metadata, en_text = en_entry

            report.write(f"Diferença {diff_index}:\n")
            if pt_type == "parágrafo":
                report.write(f"Tipo: Parágrafo\nNúmero da linha no documento: {pt_metadata[0]}\n")
            else:
                report.write(f"Tipo: Tabela\nNúmero da tabela: {pt_metadata[0]}\nLinha da tabela: {pt_metadata[1]}\nColuna da tabela: {pt_metadata[2]}\n")

            report.write(f"Português: {pt_text}\n")
            report.write(f"Inglês Original: {en_text}\n")
            report.write(f"Tradução para Inglês: {translated_line}\n")
            report.write(f"Similaridade: {similarity:.2%}\n")
            report.write("-" * 50 + "\n")
    else:
        report.write("Os textos estão alinhados e semelhantes!\n")

print(f"Relatório de diferenças salvo em: {diff_report_path}")
