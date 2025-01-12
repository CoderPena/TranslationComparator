from docx import Document
from deep_translator import GoogleTranslator
from difflib import SequenceMatcher
from tqdm import tqdm  # Biblioteca para barra de progresso

# Função para carregar texto de um arquivo DOCX com barra de progresso
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = []

    # Processar parágrafos com barra de progresso
    paragraphs = doc.paragraphs
    print(f"Lendo parágrafos do arquivo: {file_path}")
    for paragraph in tqdm(paragraphs, desc="Parágrafos processados"):
        if paragraph.text.strip():  # Ignorar linhas vazias
            text.append(paragraph.text.strip())

    # Processar tabelas com barra de progresso
    tables = doc.tables
    print(f"Lendo tabelas do arquivo: {file_path}")
    for table in tqdm(tables, desc="Tabelas processadas"):
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:  # Ignorar células vazias
                    text.append(cell_text)

    return text

# Função para contar palavras em uma lista de texto
def count_words(text_list):
    return sum(len(line.split()) for line in text_list)

# Função para comparar textos com barra de progresso
def compare_texts(pt_text, en_text):
    differences = []
    print("Comparando textos...")
    for pt_line, en_line in tqdm(zip(pt_text, en_text), total=min(len(pt_text), len(en_text)), desc="Comparações realizadas"):
        translated_line = GoogleTranslator(source='portuguese', target='english').translate(pt_line)
        similarity = SequenceMatcher(None, translated_line, en_line).ratio()
        if similarity < 0.9:  # Ajuste o limiar de similaridade
            differences.append((pt_line, en_line, translated_line, similarity))
    return differences

# Caminhos dos arquivos
pt_file = "/home/pena/Downloads/Tema_9_-_ESPERANÇA2024-12-06_16_56_37-1.docx"
en_file = "/home/pena/Downloads/Tema_9_-_ESPERANÇA2024-12-06_16_56_37-1 PTBR-EN.docx"

# Carregar os textos dos arquivos
pt_text = extract_text_from_docx(pt_file)
en_text = extract_text_from_docx(en_file)

# Contar palavras nos textos
pt_word_count = count_words(pt_text)
en_word_count = count_words(en_text)

# Comparar os textos
differences = compare_texts(pt_text, en_text)

# Gerar relatório
diff_report_path = "/home/pena/Downloads/comparacao_traducao_relatorio.txt"
with open(diff_report_path, "w", encoding="utf-8") as report:
    report.write(f"Arquivo em Português: {pt_file}\n")
    report.write(f"Palavras lidas: {pt_word_count}\n\n")
    report.write(f"Arquivo em Inglês: {en_file}\n")
    report.write(f"Palavras lidas: {en_word_count}\n\n")

    if differences:
        report.write("Diferenças encontradas:\n\n")
        for pt, en, translated, similarity in differences:
            report.write(f"Português: {pt}\n")
            report.write(f"Inglês Original: {en}\n")
            report.write(f"Tradução para Inglês: {translated}\n")
            report.write(f"Similaridade: {similarity:.2%}\n")
            report.write("-" * 50 + "\n")
    else:
        report.write("Os textos estão alinhados e semelhantes!\n")

print(f"Relatório de diferenças salvo em: {diff_report_path}")
